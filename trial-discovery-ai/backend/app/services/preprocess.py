import json
import re
from datetime import datetime, timezone
from hashlib import sha256
from io import BytesIO
from typing import Any

import fitz
import mailparser
import pytesseract
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image
from langdetect import LangDetectException, detect
from sqlalchemy.orm import Session

from app.db.models.artifact import Artifact
from app.db.models.document import Document
from app.db.models.enums import ArtifactKind, DocumentStatus
from app.storage.dropbox import DropboxClient
from app.storage.s3 import S3Client

try:
    import pillow_heif  # type: ignore[import-not-found]

    pillow_heif.register_heif_opener()
except Exception:  # noqa: BLE001
    pillow_heif = None


class PreprocessResult:
    def __init__(self, artifacts_created: int, ocr_used: bool) -> None:
        self.artifacts_created = artifacts_created
        self.ocr_used = ocr_used


def _build_base_meta(document: Document) -> dict[str, Any]:
    return {
        "source_path": document.source_path,
        "original_filename": document.original_filename,
        "mime_type": document.mime_type,
        "sha256": document.sha256,
        "file_size": document.file_size,
        "ingested_at": document.ingested_at.isoformat()
        if document.ingested_at
        else None,
    }


def _serialize_meta_value(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, (str, int, float, bool)):
        return value
    return str(value)


def _extract_properties(obj: Any, fields: list[str]) -> dict[str, Any]:
    data: dict[str, Any] = {}
    for field in fields:
        value = getattr(obj, field, None)
        if value in (None, ""):
            continue
        data[field] = _serialize_meta_value(value)
    return data


def _format_addresses(addresses: Any) -> list[str]:
    if not addresses:
        return []
    formatted: list[str] = []
    for item in addresses:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            name, email = item[0], item[1]
            if name:
                formatted.append(f"{name} <{email}>")
            else:
                formatted.append(str(email))
        else:
            formatted.append(str(item))
    return formatted


def _normalize_text(text: str) -> str:
    cleaned = text.lower()
    cleaned = re.sub(r"[^\w\s]", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _detect_language(text: str) -> str | None:
    sample = text.strip()
    if not sample:
        return None
    sample = sample[:5000]
    try:
        return detect(sample)
    except LangDetectException:
        return None


def _compute_simhash(text: str, bits: int = 64) -> str | None:
    tokens = re.findall(r"\w+", text.lower())
    if not tokens:
        return None
    vector = [0] * bits
    for token in tokens:
        token_hash = int(sha256(token.encode("utf-8")).hexdigest(), 16)
        for i in range(bits):
            bit = (token_hash >> i) & 1
            vector[i] += 1 if bit else -1
    fingerprint = 0
    for i, value in enumerate(vector):
        if value > 0:
            fingerprint |= 1 << i
    return f"{fingerprint:016x}"


def _privilege_signals(text: str) -> dict[str, Any]:
    if not text:
        return {"terms": [], "matches": 0}
    terms = [
        "attorney-client",
        "privileged",
        "work product",
        "legal advice",
        "counsel",
        "defense strategy",
        "trial strategy",
    ]
    lowered = text.lower()
    matches = [term for term in terms if term in lowered]
    return {"terms": matches, "matches": len(matches)}


def _build_text_stats(pages: list[dict[str, Any]]) -> dict[str, Any]:
    full_text = "\n\n".join(page.get("text", "") for page in pages)
    normalized = _normalize_text(full_text)
    ocr_confidences = [
        page.get("ocr_confidence")
        for page in pages
        if page.get("ocr_used") and page.get("ocr_confidence") is not None
    ]
    avg_conf = None
    if ocr_confidences:
        avg_conf = sum(ocr_confidences) / len(ocr_confidences)
    return {
        "char_count": len(full_text),
        "word_count": len(re.findall(r"\w+", full_text)),
        "normalized_sha256": sha256(normalized.encode("utf-8")).hexdigest()
        if normalized
        else None,
        "simhash": _compute_simhash(normalized) if normalized else None,
        "language": _detect_language(full_text),
        "ocr_confidence_avg": avg_conf,
        "privilege_signals": _privilege_signals(full_text),
    }


def preprocess_document(*, session: Session, document_id: str) -> PreprocessResult:
    document = session.get(Document, document_id)
    if not document:
        raise ValueError("Document not found")

    s3 = S3Client()
    if document.source_path.startswith("s3://"):
        raw_bytes = _load_s3_bytes(s3, document.source_path)
    else:
        dropbox = DropboxClient()
        raw_bytes = dropbox.download(document.source_path)

    if document.mime_type in {"application/pdf"} or document.original_filename.lower().endswith(
        ".pdf"
    ):
        return _process_pdf(
            session=session, document=document, s3=s3, raw_bytes=raw_bytes
        )

    if document.original_filename.lower().endswith(".eml"):
        return _process_eml(
            session=session, document=document, s3=s3, raw_bytes=raw_bytes
        )

    if document.original_filename.lower().endswith(".docx"):
        return _process_docx(
            session=session, document=document, s3=s3, raw_bytes=raw_bytes
        )

    if document.original_filename.lower().endswith(".xlsx"):
        return _process_xlsx(
            session=session, document=document, s3=s3, raw_bytes=raw_bytes
        )

    if document.original_filename.lower().endswith(
        (".jpg", ".jpeg", ".png", ".tif", ".tiff", ".heic", ".heif")
    ):
        return _process_image(
            session=session, document=document, s3=s3, raw_bytes=raw_bytes
        )

    if document.original_filename.lower().endswith(".textclipping"):
        return _process_text_clipping(
            session=session, document=document, s3=s3, raw_bytes=raw_bytes
        )

    document.status = DocumentStatus.ERROR
    session.commit()
    raise ValueError("Unsupported document type for preprocess")


def _process_pdf(
    *, session: Session, document: Document, s3: S3Client, raw_bytes: bytes
) -> PreprocessResult:
    pdf = fitz.open(stream=raw_bytes, filetype="pdf")
    meta = _build_base_meta(document)
    meta["document_type"] = "pdf"
    pdf_metadata = {k: _serialize_meta_value(v) for k, v in pdf.metadata.items() if v}
    if pdf_metadata:
        meta["pdf_metadata"] = pdf_metadata
    meta["page_count"] = pdf.page_count
    pages = []
    artifacts_created = 0
    ocr_used_any = False

    for page_index in range(pdf.page_count):
        page = pdf.load_page(page_index)
        text = page.get_text().strip()
        ocr_used = False

        ocr_confidence = None
        if not text:
            pix = page.get_pixmap(dpi=200)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(image).strip()
            try:
                ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [
                    float(conf)
                    for conf in ocr_data.get("conf", [])
                    if conf not in (None, "-1", -1)
                ]
                if confidences:
                    ocr_confidence = sum(confidences) / len(confidences)
            except Exception:  # noqa: BLE001
                ocr_confidence = None
            ocr_used = True
            ocr_used_any = True

            image_bytes = BytesIO()
            image.save(image_bytes, format="PNG")
            image_ref, image_hash = _store_artifact(
                s3=s3,
                document=document,
                kind=ArtifactKind.PAGE_IMAGE,
                data=image_bytes.getvalue(),
                content_type="image/png",
                suffix=f"page_{page_index + 1}.png",
            )
            _create_artifact_row(
                session, document.id, image_ref, image_hash, ArtifactKind.PAGE_IMAGE
            )
            artifacts_created += 1

        pages.append(
            {
                "page_num": page_index + 1,
                "text": text,
                "ocr_used": ocr_used,
                "ocr_confidence": ocr_confidence,
            }
        )

    extracted_payload = {
        "document_id": str(document.id),
        "page_count": pdf.page_count,
        "meta": meta,
        "pages": pages,
    }

    meta["text_stats"] = _build_text_stats(pages)

    extracted_bytes = json.dumps(extracted_payload, ensure_ascii=False).encode("utf-8")
    extracted_ref, extracted_hash = _store_artifact(
        s3=s3,
        document=document,
        kind=ArtifactKind.EXTRACTED_TEXT,
        data=extracted_bytes,
        content_type="application/json",
        suffix="extracted_text.json",
    )
    _create_artifact_row(
        session,
        document.id,
        extracted_ref,
        extracted_hash,
        ArtifactKind.EXTRACTED_TEXT,
    )
    artifacts_created += 1

    document.page_count = pdf.page_count
    document.status = DocumentStatus.PREPROCESSED
    session.commit()
    return PreprocessResult(artifacts_created=artifacts_created, ocr_used=ocr_used_any)


def _process_eml(
    *, session: Session, document: Document, s3: S3Client, raw_bytes: bytes
) -> PreprocessResult:
    mail = mailparser.parse_from_bytes(raw_bytes)

    if mail.text_plain:
        body_text = "\n".join(mail.text_plain)
    else:
        body_text = mail.body or ""

    meta = _build_base_meta(document)
    meta["document_type"] = "email"
    meta["email"] = {
        "subject": mail.subject,
        "from": _format_addresses(mail.from_),
        "to": _format_addresses(mail.to),
        "cc": _format_addresses(mail.cc),
        "bcc": _format_addresses(mail.bcc),
        "date": _serialize_meta_value(mail.date),
        "message_id": mail.message_id,
    }

    payload: dict[str, Any] = {
        "subject": mail.subject,
        "from": mail.from_,
        "to": mail.to,
        "cc": mail.cc,
        "bcc": mail.bcc,
        "date": mail.date,
        "message_id": mail.message_id,
        "body": body_text,
        "attachments": [a.get("filename") for a in mail.attachments],
    }

    extracted_payload = {
        "document_id": str(document.id),
        "page_count": 1,
        "meta": meta,
        "pages": [
            {
                "page_num": 1,
                "text": body_text,
                "ocr_used": False,
                "ocr_confidence": None,
            }
        ],
    }

    meta["text_stats"] = _build_text_stats(extracted_payload["pages"])

    email_bytes = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    email_ref, email_hash = _store_artifact(
        s3=s3,
        document=document,
        kind=ArtifactKind.EMAIL_JSON,
        data=email_bytes,
        content_type="application/json",
        suffix="email.json",
    )
    _create_artifact_row(
        session, document.id, email_ref, email_hash, ArtifactKind.EMAIL_JSON
    )

    extracted_bytes = json.dumps(extracted_payload, ensure_ascii=False).encode("utf-8")
    extracted_ref, extracted_hash = _store_artifact(
        s3=s3,
        document=document,
        kind=ArtifactKind.EXTRACTED_TEXT,
        data=extracted_bytes,
        content_type="application/json",
        suffix="extracted_text.json",
    )
    _create_artifact_row(
        session,
        document.id,
        extracted_ref,
        extracted_hash,
        ArtifactKind.EXTRACTED_TEXT,
    )

    document.page_count = 1
    document.status = DocumentStatus.PREPROCESSED
    session.commit()

    return PreprocessResult(artifacts_created=2, ocr_used=False)


def _process_docx(
    *, session: Session, document: Document, s3: S3Client, raw_bytes: bytes
) -> PreprocessResult:
    doc = DocxDocument(BytesIO(raw_bytes))
    parts: list[str] = []
    meta = _build_base_meta(document)
    meta["document_type"] = "docx"
    meta["docx_properties"] = _extract_properties(
        doc.core_properties,
        [
            "author",
            "category",
            "comments",
            "content_status",
            "created",
            "identifier",
            "keywords",
            "language",
            "last_modified_by",
            "last_printed",
            "modified",
            "revision",
            "subject",
            "title",
            "version",
        ],
    )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    text = "\n".join(parts).strip()

    extracted_payload = {
        "document_id": str(document.id),
        "page_count": 1,
        "meta": meta,
        "pages": [
            {
                "page_num": 1,
                "text": text,
                "ocr_used": False,
                "ocr_confidence": None,
            }
        ],
    }

    meta["text_stats"] = _build_text_stats(extracted_payload["pages"])

    extracted_bytes = json.dumps(extracted_payload, ensure_ascii=False).encode("utf-8")
    extracted_ref, extracted_hash = _store_artifact(
        s3=s3,
        document=document,
        kind=ArtifactKind.EXTRACTED_TEXT,
        data=extracted_bytes,
        content_type="application/json",
        suffix="extracted_text.json",
    )
    _create_artifact_row(
        session,
        document.id,
        extracted_ref,
        extracted_hash,
        ArtifactKind.EXTRACTED_TEXT,
    )

    document.page_count = 1
    document.status = DocumentStatus.PREPROCESSED
    session.commit()
    return PreprocessResult(artifacts_created=1, ocr_used=False)


def _process_xlsx(
    *, session: Session, document: Document, s3: S3Client, raw_bytes: bytes
) -> PreprocessResult:
    workbook = load_workbook(filename=BytesIO(raw_bytes), data_only=True)
    lines: list[str] = []
    meta = _build_base_meta(document)
    meta["document_type"] = "xlsx"
    meta["xlsx_properties"] = _extract_properties(
        workbook.properties,
        [
            "creator",
            "lastModifiedBy",
            "created",
            "modified",
            "title",
            "subject",
            "description",
            "keywords",
            "category",
            "contentStatus",
            "identifier",
            "language",
            "revision",
            "version",
        ],
    )

    for sheet in workbook.worksheets:
        lines.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell) for cell in row]
            if any(cell.strip() for cell in cells):
                lines.append("\t".join(cells))
        lines.append("")

    text = "\n".join(lines).strip()

    extracted_payload = {
        "document_id": str(document.id),
        "page_count": 1,
        "meta": meta,
        "pages": [
            {
                "page_num": 1,
                "text": text,
                "ocr_used": False,
                "ocr_confidence": None,
            }
        ],
    }

    meta["text_stats"] = _build_text_stats(extracted_payload["pages"])

    extracted_bytes = json.dumps(extracted_payload, ensure_ascii=False).encode("utf-8")
    extracted_ref, extracted_hash = _store_artifact(
        s3=s3,
        document=document,
        kind=ArtifactKind.EXTRACTED_TEXT,
        data=extracted_bytes,
        content_type="application/json",
        suffix="extracted_text.json",
    )
    _create_artifact_row(
        session,
        document.id,
        extracted_ref,
        extracted_hash,
        ArtifactKind.EXTRACTED_TEXT,
    )

    document.page_count = 1
    document.status = DocumentStatus.PREPROCESSED
    session.commit()
    return PreprocessResult(artifacts_created=1, ocr_used=False)


def _process_image(
    *, session: Session, document: Document, s3: S3Client, raw_bytes: bytes
) -> PreprocessResult:
    if document.original_filename.lower().endswith((".heic", ".heif")) and pillow_heif is None:
        document.status = DocumentStatus.ERROR
        session.commit()
        raise ValueError("HEIC/HEIF support requires pillow-heif")

    image = Image.open(BytesIO(raw_bytes))
    text = pytesseract.image_to_string(image).strip()
    ocr_confidence = None
    try:
        ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        confidences = [
            float(conf)
            for conf in ocr_data.get("conf", [])
            if conf not in (None, "-1", -1)
        ]
        if confidences:
            ocr_confidence = sum(confidences) / len(confidences)
    except Exception:  # noqa: BLE001
        ocr_confidence = None
    meta = _build_base_meta(document)
    meta["document_type"] = "image"
    meta["image"] = {
        "format": image.format,
        "mode": image.mode,
        "size": list(image.size),
    }

    extracted_payload = {
        "document_id": str(document.id),
        "page_count": 1,
        "meta": meta,
        "pages": [
            {
                "page_num": 1,
                "text": text,
                "ocr_used": True,
                "ocr_confidence": ocr_confidence,
            }
        ],
    }

    meta["text_stats"] = _build_text_stats(extracted_payload["pages"])

    extracted_bytes = json.dumps(extracted_payload, ensure_ascii=False).encode("utf-8")
    extracted_ref, extracted_hash = _store_artifact(
        s3=s3,
        document=document,
        kind=ArtifactKind.EXTRACTED_TEXT,
        data=extracted_bytes,
        content_type="application/json",
        suffix="extracted_text.json",
    )
    _create_artifact_row(
        session,
        document.id,
        extracted_ref,
        extracted_hash,
        ArtifactKind.EXTRACTED_TEXT,
    )

    document.page_count = 1
    document.status = DocumentStatus.PREPROCESSED
    session.commit()
    return PreprocessResult(artifacts_created=1, ocr_used=True)


def _process_text_clipping(
    *, session: Session, document: Document, s3: S3Client, raw_bytes: bytes
) -> PreprocessResult:
    text = raw_bytes.decode("utf-8", errors="replace").strip()
    meta = _build_base_meta(document)
    meta["document_type"] = "textclipping"

    extracted_payload = {
        "document_id": str(document.id),
        "page_count": 1,
        "meta": meta,
        "pages": [
            {
                "page_num": 1,
                "text": text,
                "ocr_used": False,
                "ocr_confidence": None,
            }
        ],
    }

    meta["text_stats"] = _build_text_stats(extracted_payload["pages"])

    extracted_bytes = json.dumps(extracted_payload, ensure_ascii=False).encode("utf-8")
    extracted_ref, extracted_hash = _store_artifact(
        s3=s3,
        document=document,
        kind=ArtifactKind.EXTRACTED_TEXT,
        data=extracted_bytes,
        content_type="application/json",
        suffix="extracted_text.json",
    )
    _create_artifact_row(
        session,
        document.id,
        extracted_ref,
        extracted_hash,
        ArtifactKind.EXTRACTED_TEXT,
    )

    document.page_count = 1
    document.status = DocumentStatus.PREPROCESSED
    session.commit()
    return PreprocessResult(artifacts_created=1, ocr_used=False)


def _store_artifact(
    *,
    s3: S3Client,
    document: Document,
    kind: ArtifactKind,
    data: bytes,
    content_type: str,
    suffix: str,
):
    key = f"matters/{document.matter_id}/documents/{document.id}/{kind.value}/{suffix}"
    content_hash = sha256(data).hexdigest()
    obj_ref = s3.put_bytes(key=key, data=data, content_type=content_type)
    return obj_ref, content_hash


def _create_artifact_row(
    session: Session,
    document_id,
    obj_ref,
    content_hash: str,
    kind: ArtifactKind,
) -> None:
    artifact = Artifact(
        document_id=document_id,
        kind=kind,
        uri=obj_ref.uri,
        content_hash=content_hash,
        created_at=datetime.now(timezone.utc),
    )
    session.add(artifact)
    session.flush()


def _load_s3_bytes(s3: S3Client, uri: str) -> bytes:
    if not uri.startswith("s3://"):
        raise ValueError("Invalid S3 URI")
    parts = uri.replace("s3://", "", 1).split("/", 1)
    bucket = parts[0]
    key = parts[1]
    return s3.get_bytes(bucket=bucket, key=key)
